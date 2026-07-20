"""Studio de rapports : génération assistée, conversion de réponses en blocs,
et export Word (.docx) / PDF / Markdown.

Hors-ligne : la génération s'appuie sur le pipeline d'analyse (agent approfondi)
quand une source est fournie — le rapport est donc chiffré et argumenté, jamais
inventé. Sans source, on produit un plan à compléter.
"""
from __future__ import annotations

import io
import re


# ---------------------------------------------------------------------------
# Conversion d'une réponse de chat en blocs de rapport
# ---------------------------------------------------------------------------
def response_to_blocks(title: str, resp: dict) -> list[dict]:
    """Transforme une réponse d'analyse (ChatResponse sérialisée) en blocs."""
    blocks: list[dict] = []
    md_lines: list[str] = [f"## {title}"]

    analysis = resp.get("analysis") or {}
    if analysis.get("summary"):
        md_lines.append(analysis["summary"])
    for o in analysis.get("observations", []) or []:
        md_lines.append(f"- {o}")

    deep = resp.get("deep")
    if deep:
        for c in deep.get("context", []) or []:
            md_lines.append(c)
        if deep.get("drivers"):
            md_lines.append("### Facteurs explicatifs")
            md_lines += [f"- {d}" for d in deep["drivers"]]
        if deep.get("findings"):
            md_lines.append("### Points d'attention")
            md_lines += [f"- {f}" for f in deep["findings"]]
        if deep.get("recommendations"):
            md_lines.append("### Recommandations")
            md_lines += [f"- {r}" for r in deep["recommendations"]]
    elif analysis.get("recommendations"):
        md_lines.append("### Recommandations")
        md_lines += [f"- {r}" for r in analysis["recommendations"]]

    blocks.append({"kind": "markdown", "content": {"text": "\n\n".join(md_lines)}})

    # Croisement -> tableau
    if deep and deep.get("crosstab") and deep["crosstab"].get("cells"):
        ct = deep["crosstab"]
        cols = [ct["dim_a"], ct["dim_b"], ct.get("metric", "valeur")]
        rows = [[c["a"], c["b"], c["value"]] for c in ct["cells"]]
        blocks.append({"kind": "table", "content": {
            "caption": f"Croisement {ct['dim_a']} × {ct['dim_b']}", "columns": cols, "rows": rows,
        }})

    # Graphique (données + suggestion) si graphable
    chart = resp.get("chart")
    if chart and chart.get("type") not in (None, "table") and resp.get("columns"):
        blocks.append({"kind": "chart", "content": {
            "caption": resp.get("question") or title,
            "columns": resp["columns"], "rows": resp.get("rows", [])[:200], "chart": chart,
        }})

    # Données brutes (aperçu)
    if resp.get("columns") and resp.get("rows"):
        blocks.append({"kind": "table", "content": {
            "caption": "Données",
            "columns": resp["columns"], "rows": resp["rows"][:50],
        }})
    return blocks


def skeleton_blocks(prompt: str) -> list[dict]:
    """Sans source : un plan de rapport à compléter (l'IA cloud le remplirait)."""
    text = (
        f"# {prompt.strip()}\n\n"
        "## Contexte\n\n_À compléter._\n\n"
        "## Analyse\n\n_À compléter (connectez une source pour une analyse chiffrée)._\n\n"
        "## Recommandations\n\n_À compléter._"
    )
    return [{"kind": "markdown", "content": {"text": text}}]


# ---------------------------------------------------------------------------
# Rendu Markdown
# ---------------------------------------------------------------------------
def _table_md(columns: list, rows: list) -> str:
    head = "| " + " | ".join(str(c) for c in columns) + " |"
    sep = "| " + " | ".join("---" for _ in columns) + " |"
    body = "\n".join("| " + " | ".join(str(v) for v in r) + " |" for r in rows)
    return f"{head}\n{sep}\n{body}"


def to_markdown(title: str, blocks: list[dict]) -> str:
    out = [f"# {title}", ""]
    for b in blocks:
        c = b.get("content") or {}
        if b["kind"] == "markdown":
            out.append(c.get("text", ""))
        elif b["kind"] == "table":
            if c.get("caption"):
                out.append(f"**{c['caption']}**")
            out.append(_table_md(c.get("columns", []), c.get("rows", [])))
        elif b["kind"] == "chart":
            cap = c.get("caption", "Graphique")
            out.append(f"**{cap}** _(graphique {c.get('chart', {}).get('type', '')})_")
            out.append(_table_md(c.get("columns", []), c.get("rows", [])[:20]))
        out.append("")
    return "\n".join(out)


# ---------------------------------------------------------------------------
# Export Word (.docx)
# ---------------------------------------------------------------------------
def _emit_markdown_docx(doc, text: str) -> None:
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line.replace("**", ""))


def to_docx(title: str, blocks: list[dict]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for b in blocks:
        c = b.get("content") or {}
        if b["kind"] == "markdown":
            _emit_markdown_docx(doc, c.get("text", ""))
        elif b["kind"] in ("table", "chart"):
            if c.get("caption"):
                doc.add_heading(c["caption"], level=3)
            if b["kind"] == "chart":
                doc.add_paragraph(f"(graphique {c.get('chart', {}).get('type', '')})")
            columns = c.get("columns", [])
            rows = c.get("rows", [])[:60]
            if columns:
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = "Light Grid Accent 1"
                for i, col in enumerate(columns):
                    table.rows[0].cells[i].text = str(col)
                for r in rows:
                    cells = table.add_row().cells
                    for i, v in enumerate(r):
                        cells[i].text = "" if v is None else str(v)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Export PDF
# ---------------------------------------------------------------------------
_TYPO = {"—": "-", "–": "-", "’": "'", "‘": "'",
         "“": '"', "”": '"', "€": "EUR", "→": "->",
         "×": "x", "≥": ">=", "≤": "<=", "…": "..."}


def _latin1(s: str) -> str:
    for k, v in _TYPO.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")


def to_pdf(title: str, blocks: list[dict]) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    # new_x/new_y : ramène le curseur à la marge gauche (sinon fpdf2 le laisse à
    # droite et la ligne suivante n'a « plus de place »).
    pdf.multi_cell(0, 10, _latin1(title), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    def para(text: str, size=11, style=""):
        pdf.set_font("Helvetica", style, size)
        pdf.multi_cell(0, 6, _latin1(text), new_x="LMARGIN", new_y="NEXT")

    def table(columns, rows):
        pdf.set_font("Helvetica", "B", 9)
        n = max(1, len(columns))
        w = (pdf.w - 20) / n
        for col in columns:
            pdf.cell(w, 6, _latin1(str(col))[:22], border=1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 9)
        for r in rows[:40]:
            for v in r:
                pdf.cell(w, 6, _latin1("" if v is None else str(v))[:22], border=1)
            pdf.ln()

    for b in blocks:
        c = b.get("content") or {}
        if b["kind"] == "markdown":
            for raw in c.get("text", "").split("\n"):
                line = raw.rstrip()
                if not line:
                    continue
                if line.startswith("### "):
                    para(line[4:], 12, "B")
                elif line.startswith("## "):
                    para(line[3:], 14, "B")
                elif line.startswith("# "):
                    para(line[2:], 16, "B")
                elif line.startswith(("- ", "* ")):
                    para("  • " + line[2:])
                else:
                    para(line.replace("**", ""))
            pdf.ln(1)
        elif b["kind"] in ("table", "chart"):
            if c.get("caption"):
                para(c["caption"], 12, "B")
            if b["kind"] == "chart":
                para(f"(graphique {c.get('chart', {}).get('type', '')})", 9, "I")
            table(c.get("columns", []), c.get("rows", []))
            pdf.ln(2)

    out = pdf.output()
    return bytes(out)


# ---------------------------------------------------------------------------
def default_title(prompt: str) -> str:
    t = re.sub(r"\s+", " ", prompt).strip()
    return (t[:60] + "…") if len(t) > 60 else (t or "Nouveau rapport")
